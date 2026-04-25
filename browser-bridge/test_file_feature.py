#!/usr/bin/env python3
"""Test sage.py file attachment feature."""
import sys, os, io
sys.path.insert(0, '/opt/browser-bridge')
import sage
from sage import MAX_FILE_CHARS, extract_text_from_file, supported_ext, load_file_context

PASS = "OK"
FAIL = "FAIL"
results = []

def check(name, cond):
    results.append((name, PASS if cond else FAIL))
    print(f"  [{PASS if cond else FAIL}] {name}")

print("=== 1. supported_ext ===")
check("md supported", supported_ext("readme.md"))
check("py supported", supported_ext("script.py"))
check("csv supported", supported_ext("data.csv"))
check("pdf supported", supported_ext("doc.pdf"))
check("xlsx supported", supported_ext("data.xlsx"))
check("jpg NOT supported", not supported_ext("image.jpg"))
check("zip NOT supported", not supported_ext("archive.zip"))
check("exe NOT supported", not supported_ext("virus.exe"))

print()
print("=== 2. extract_text_from_file ===")

# txt utf-8
data = "Hello World\nLine 2".encode("utf-8")
t = extract_text_from_file(data, "test.txt")
check("txt utf-8", "Hello World" in t)

# txt cp1251
data = "Windows text".encode("cp1251")
t = extract_text_from_file(data, "win.txt")
check("txt cp1251", "Windows text" in t)

# json
data = '{"key": "value"}'.encode("utf-8")
t = extract_text_from_file(data, "cfg.json")
check("json", "value" in t)

# python code
data = b"def hello():\n    return 42"
t = extract_text_from_file(data, "code.py")
check("python code", "def hello" in t)

# csv
data = "Name,Age\nIvan,25\nMaria,30".encode("utf-8")
t = extract_text_from_file(data, "data.csv")
check("csv", "Ivan" in t and "Maria" in t)

# sql
data = b"SELECT * FROM users WHERE id = 1;"
t = extract_text_from_file(data, "query.sql")
check("sql", "SELECT" in t)

# yaml
data = b"key: value\nlist:\n  - item1\n  - item2"
t = extract_text_from_file(data, "config.yaml")
check("yaml", "item1" in t)

# unsupported format
try:
    extract_text_from_file(b"binary data", "file.exe")
    check("unsupported raises ValueError", False)
except ValueError:
    check("unsupported raises ValueError", True)

print()
print("=== 3. DOCX ===")
try:
    import docx as _docx
    doc = _docx.Document()
    doc.add_paragraph("Test document content")
    doc.add_paragraph("Second paragraph")
    buf = io.BytesIO()
    doc.save(buf)
    t = extract_text_from_file(buf.getvalue(), "test.docx")
    check("docx content parsed", "Test document content" in t)
    check("docx second paragraph", "Second paragraph" in t)
except Exception as e:
    check("docx (ERROR: " + str(e) + ")", False)

print()
print("=== 4. XLSX ===")
try:
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws.append(["Product", "Price", "Stock"])
    ws.append(["Table", 1500, 10])
    ws.append(["Chair", 800, 25])
    buf = io.BytesIO()
    wb.save(buf)
    t = extract_text_from_file(buf.getvalue(), "test.xlsx")
    check("xlsx header row", "Product" in t)
    check("xlsx data row", "Table" in t and "Chair" in t)
    check("xlsx sheet name", "Sheet1" in t)
except Exception as e:
    check("xlsx (ERROR: " + str(e) + ")", False)

print()
print("=== 5. Truncation (MAX_FILE_CHARS=" + str(MAX_FILE_CHARS) + ") ===")
original_fetch = sage.fetch_file
big_data = ("B" * (MAX_FILE_CHARS + 1000)).encode("utf-8")
sage.fetch_file = lambda url: (big_data, "big.txt")
text, name = load_file_context("http://fake/big.txt")
sage.fetch_file = original_fetch
check("truncated to max", len(text) <= MAX_FILE_CHARS + 100)
check("truncation marker added", "обрезан" in text)

print()
print("=== 6. cmd_ask with file ===")
calls = []

def mock_bridge(path, payload):
    calls.append(payload)
    return {"responses": [{"llm": "deepseek", "text": "Mock answer from LLM for file test"}]}

def mock_cerebras(msgs, max_tokens=700):
    return "Synthesis result"

original_bridge = sage.bridge_post
original_cerebras = sage.cerebras_chat
sage.bridge_post = mock_bridge
sage.cerebras_chat = mock_cerebras
sage.fetch_file = lambda url: ("File content here".encode("utf-8"), "doc.txt")

buf = io.StringIO()
from contextlib import redirect_stdout
with redirect_stdout(buf):
    sage.cmd_ask(99999, "Analyze this", file_url="http://fake/doc.txt", file_name="doc.txt")
output = buf.getvalue()

sage.bridge_post = original_bridge
sage.cerebras_chat = original_cerebras
sage.fetch_file = original_fetch

msg_sent = calls[0]["message"] if calls else ""
check("file content in LLM message", "File content here" in msg_sent)
check("question in LLM message", "Analyze this" in msg_sent)
check("filename in output", "doc.txt" in output)
check("file emoji in output", chr(128206) in output)  # 📎
check("SESSION_ID in output", "SESSION_ID:" in output)
check("synthesis in output", "Synthesis result" in output)
check("separator in LLM message", "---" in msg_sent)

print()
print("=== 7. cmd_ask without file (backward compat) ===")
calls2 = []
sage.bridge_post = lambda path, payload: (calls2.append(payload), {"responses": [{"llm": "deepseek", "text": "Normal answer"}]})[1]
sage.cerebras_chat = lambda msgs, max_tokens=700: "Normal synthesis"

buf2 = io.StringIO()
with redirect_stdout(buf2):
    sage.cmd_ask(99999, "Simple question without file")
out2 = buf2.getvalue()

sage.bridge_post = original_bridge
sage.cerebras_chat = original_cerebras

msg2 = calls2[0]["message"] if calls2 else ""
check("no file: question sent", "Simple question without file" in msg2)
check("no file: no separator", "---" not in msg2)
check("no file: no emoji in header", "doc.txt" not in out2)

print()
print("=== 8. ask_file DB entry stores filename prefix ===")
calls3 = []
sage.bridge_post = lambda path, payload: (calls3.append(payload), {"responses": [{"llm": "deepseek", "text": "DB test answer"}]})[1]
sage.cerebras_chat = lambda msgs, max_tokens=700: "DB test synthesis"
sage.fetch_file = lambda url: ("DB test content".encode("utf-8"), "report.txt")

buf3 = io.StringIO()
with redirect_stdout(buf3):
    sage.cmd_ask(88888, "What is in the report", file_url="http://fake/report.txt", file_name="report.txt")
out3 = buf3.getvalue()
sage.bridge_post = original_bridge
sage.cerebras_chat = original_cerebras
sage.fetch_file = original_fetch

# Check DB entry
import sqlite3
conn = sqlite3.connect("/opt/openclaw-sage/sage.db")
conn.row_factory = sqlite3.Row
row = conn.execute(
    "SELECT question FROM messages ORDER BY id DESC LIMIT 1"
).fetchone()
conn.close()
if row:
    saved_q = row["question"]
    check("DB saves filename prefix", "[" + chr(128206) in saved_q or "report.txt" in saved_q)
    check("DB question readable (not full file content)", "DB test content" not in saved_q)
else:
    check("DB entry created", False)

print()
total = len(results)
passed = sum(1 for _, s in results if s == PASS)
failed = total - passed
print(f"=== ИТОГО: {passed}/{total} прошло, {failed} упало ===")
if failed > 0:
    print("Упавшие:")
    for name, status in results:
        if status == FAIL:
            print(f"  FAIL: {name}")
sys.exit(0 if failed == 0 else 1)
